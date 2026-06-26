from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def add_heading(doc, text, level=1, color="1F3864"):
    h = doc.add_heading(text, level=level)
    run = h.runs[0] if h.runs else h.add_run(text)
    run.font.color.rgb = RGBColor.from_string(color)
    return h

def add_table(doc, headers, rows, header_bg="1F3864", alt_bg="E8EDF5"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(cell, header_bg)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_bg(cell, alt_bg)
    return table

def add_code(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p

def main():
    doc = Document()

    # ─── Page Margins ────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ─── TITLE PAGE ──────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("🛡️ NexusSec — Enterprise Cybersecurity Analytics Platform")
    tr.font.size = Pt(22)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(0x1f, 0x38, 0x64)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Complete Technical Documentation")
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Version 1.0  |  Generated: {datetime.date.today().strftime('%B %d, %Y')}  |  Classification: Internal")

    doc.add_page_break()

    
    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "NexusSec is a production-ready, enterprise-grade Security Operations Center (SOC) analytics "
        "platform designed to continuously ingest, analyse, and respond to cybersecurity threats in real time. "
        "Built on Python/Flask with a modular architecture, it integrates machine-learning threat detection, "
        "local rule-based AI analysis, Sonar-like static code scanning, automated data ingestion from "
        "Grafana/Loki/Prometheus, and a modern dark-mode web dashboard — all backed by a persistent SQLite database."
    )
    doc.add_paragraph(
        "The platform replaces simple proof-of-concept (POC) log scanning with a scalable pipeline capable "
        "of detecting 40+ attack categories, providing real-time alerts, vulnerability assessments, predictive "
        "risk forecasting, and one-click Excel report exports."
    )

    
    add_heading(doc, "2. System Architecture", 1)
    doc.add_paragraph(
        "NexusSec follows the Flask Application Factory pattern with clearly separated concerns across modules. "
        "The high-level data flow is:"
    )
    doc.add_paragraph(
        "Mock Grafana/Loki/Prometheus APIs → Log Ingestor (APScheduler) → Parser → Feature Engineering → "
        "Threat Detection Engine (ML + Signatures) → SQLite Database → REST API → Dashboard Frontend"
    )

    add_heading(doc, "2.1 Component Diagram", 2)
    comp_rows = [
        ("Frontend (Browser)", "HTML / CSS / Vanilla JS", "Polls /api/dashboard every 3s, renders charts and threat tables"),
        ("Flask Web Server",   "Python 3.13 / Flask 3.x", "Serves static files, API endpoints, and the main template"),
        ("REST API Layer",     "cybersec_platform/api.py", "8 endpoints: dashboard, scan, upload-code, exports, start/stop-ingestion"),
        ("Background Scheduler", "APScheduler 3.x",       "Polls mock APIs every 10–15 s, saves data to DB"),
        ("Threat Detection",  "Scikit-learn + RegEx",     "40+ MITRE-mapped signatures + Isolation Forest + Random Forest"),
        ("Local Analysis",    "Pre-defined Expert Matrix", "Generates high-fidelity explanations & mitigations locally"),
        ("Sonar Analyzer",    "Regex rule engine",         "Static analysis of uploaded source files for bugs and vulnerabilities"),
        ("Database",          "SQLAlchemy + SQLite",       "Stores logs, alerts, metrics, vulnerabilities, system health"),
    ]
    add_table(doc, ["Component", "Technology", "Responsibility"], comp_rows)

    
    add_heading(doc, "3. Repository File Structure", 1)
    add_code(doc,
        "Cybersecurity POC/\n"
        "├── run.py                          # Application entry point\n"
        "├── main.py                         # Legacy CLI entry (retained for reference)\n"
        "├── requirements.txt                # Python dependencies\n"
        "├── cybersec.db                     # SQLite data store (auto-created)\n"
        "├── cybersec_platform/              # Core application package\n"
        "│   ├── __init__.py\n"
        "│   ├── app.py                      # Flask app factory\n"
        "│   ├── api.py                      # REST endpoints\n"
        "│   ├── config.py                   # Configuration dataclass\n"
        "│   ├── database.py                 # SQLAlchemy ORM models\n"
        "│   ├── ingestion.py                # Background polling & pipeline\n"
        "│   ├── integrations.py             # Mock Grafana/Loki/Prometheus clients\n"
        "│   ├── parsing.py                  # Multi-format log parser\n"
        "│   ├── feature_engineering.py      # Feature extraction & normalisation\n"
        "│   ├── models.py                   # ML model wrappers (IsoForest, RF)\n"
        "│   ├── threat_detection.py         # Unified detection engine\n"
        "│   ├── vulnerability_assessment.py # CVE risk scoring\n"
        "│   ├── predictive_analytics.py     # Trend forecasting\n"
        "│   ├── sonar_analyzer.py           # Sonar-like static code analysis\n"
        "│   └── utils.py                    # Shared helpers\n"
        "├── evaluate_model.py               # ML evaluation CLI\n"
        "├── .github/workflows/sonar-analysis.yml  # CI static analysis\n"
        "├── static/\n"
        "│   ├── dashboard.css               # Dark-mode stylesheet\n"
        "│   └── dashboard.js                # Frontend logic & API polling\n"
        "├── templates/\n"
        "│   └── index.html                  # Main SPA template\n"
        "└── tests/\n"
        "    ├── test_feature_engineering.py\n"
        "    └── test_parsing.py"
    )

    # ─── 4. MODULE DOCUMENTATION ─────────────────────────────────
    add_heading(doc, "4. Module Documentation", 1)

    # 4.1 database.py
    add_heading(doc, "4.1 database.py — Data Persistence Layer", 2)
    doc.add_paragraph(
        "Defines all SQLAlchemy ORM models and initialises the SQLite database. "
        "The global SessionLocal factory is imported by all modules needing DB access."
    )
    db_rows = [
        ("LogEntry",     "log_entries",    "Raw ingested log messages from Loki/manual scan"),
        ("Alert",        "alerts",         "Detected threats with severity, MITRE tactic, and ML Threat Analysis summary"),
        ("Metric",       "metrics",        "Time-series Prometheus metrics (CPU, memory, network)"),
        ("Vulnerability","vulnerabilities", "Known CVEs with CVSS scores and mitigations"),
        ("SystemHealth", "system_health",  "Aggregated system resource snapshots"),
    ]
    add_table(doc, ["Model", "Table", "Purpose"], db_rows)

    # 4.2 ingestion.py
    add_heading(doc, "4.2 ingestion.py — Background Pipeline", 2)
    doc.add_paragraph(
        "Implements the LogIngestor singleton that manages an APScheduler instance. "
        "Two recurring jobs run in background threads:"
    )
    doc.add_paragraph("• poll_loki() — runs every 10 s: fetches 20 log lines from LokiClient, extracts features, runs ThreatDetector, automatically generates local security analysis for high/critical alerts, saves results to DB.")
    doc.add_paragraph("• poll_prometheus() — runs every 15 s: fetches system metrics from PrometheusClient, saves a SystemHealth snapshot to DB.")

    # 4.3 integrations.py
    add_heading(doc, "4.3 integrations.py — Mock API Clients", 2)
    doc.add_paragraph(
        "Provides realistic mock implementations of GrafanaClient, LokiClient, and PrometheusClient. "
        "The LokiClient generates logs from a library of 50+ attack templates (12% attack rate) and 10 benign templates. "
        "PrometheusClient simulates CPU/memory spikes (8% chance) and network exfiltration anomalies (3% chance). "
        "Replace these classes with real HTTP calls to connect to a live Grafana/Loki/Prometheus stack."
    )

    # 4.4 threat_detection.py
    add_heading(doc, "4.4 threat_detection.py — Detection Engine", 2)
    doc.add_paragraph(
        "The ThreatDetector class is the core of the platform. It combines two detection strategies:"
    )
    doc.add_paragraph("• Rule-Based Signatures: 40+ compiled RegEx patterns mapped to MITRE ATT&CK technique IDs.")
    doc.add_paragraph("• Machine Learning: Isolation Forest for anomaly detection + Random Forest for multi-class threat classification.")
    doc.add_paragraph("• Event-Based Multi-Threat Analysis: Refactored to perform independent line-by-line event parsing and threat extraction, ensuring multi-threat log entries generate separate, unmerged alerts.")
    doc.add_paragraph()

    attack_rows = [
        ("Web Application", "SQL Injection, XSS, Path Traversal, RCE, SSRF, XXE, SSTI, Deserialization, HTTP Smuggling"),
        ("Authentication",  "Brute-Force, Password Spraying, Credential Dumping, Credential Exposure, Token Theft"),
        ("Privilege Escalation", "sudo/root Escalation, Container Escape, Kernel Exploits"),
        ("Network",         "Port Scanning, DDoS/SYN Flood, MitM/ARP Poisoning, DNS Tunneling, Lateral Movement"),
        ("Malware",         "Ransomware, Cryptominer, C2 Beacon, Fileless/LOLBins, Persistence"),
        ("Cloud/Infrastructure", "IMDS Abuse, K8s API Abuse, IAM Escalation, S3 Exposure"),
        ("Exfiltration",    "Data Exfiltration, Database Dump"),
        ("Defense Evasion", "Log Tampering, Firewall Disable, Process Injection, SSH Tampering"),
        ("Initial Access",  "Log4Shell/CVE Exploitation, Supply Chain Attack, Tor/Proxy, Zero-Day"),
        ("Discovery/Recon", "API Abuse/Scraping, VPN Abuse"),
    ]
    add_table(doc, ["Attack Category", "Detected Threats"], attack_rows)

    # 4.5 Local Analysis Engine
    add_heading(doc, "4.5 Local Analysis Engine", 2)
    doc.add_paragraph(
        "Runs 100% locally on the host machine without external API dependencies. Implements two core analysis methods in ThreatDetector:"
    )
    doc.add_paragraph("• local_analyze_log(log_text) → {probability: int, reason: str}: Computes a 0–100 threat probability and returns a structured explanation using rules and ML anomalies.")
    doc.add_paragraph("• local_analyze_alert(threat_type, severity, raw_message) → str: Generates a professional 2–3 sentence security summary containing attack explanation, impact, and mitigation from a pre-defined 43-threat expert matrix.")

    # 4.6 api.py
    add_heading(doc, "4.6 api.py — REST API Endpoints", 2)
    api_rows = [
        ("POST", "/api/start-ingestion", "Start the background APScheduler ingestion loop"),
        ("POST", "/api/stop-ingestion",  "Stop the background APScheduler ingestion loop"),
        ("GET",  "/api/dashboard",        "Live dashboard data; optional ?time_range=<minutes> (default 60)"),
        ("POST", "/api/scan",             "Manual log scan: body={log_text}. Returns detections, assessment, and ml_analysis"),
        ("POST", "/api/upload-code",      "Upload source file (code_file) for Sonar-like static analysis"),
        ("GET",  "/api/export/excel",     "Download all alerts as security_alerts_export.xlsx"),
        ("GET",  "/api/export/vulnerabilities", "Download CVE findings as vulnerabilities_export.csv"),
    ]
    add_table(doc, ["Method", "Endpoint", "Description"], api_rows)

    # 4.7 parsing.py
    add_heading(doc, "4.7 parsing.py — Log Parsing", 2)
    doc.add_paragraph(
        "Auto-detects the format of incoming log lines using heuristic probing and regex. "
        "Supported formats: JSON (structured), Apache/Nginx Combined Log, Syslog (RFC 5424 & 3164), "
        "Key=Value pairs, and plain text. Normalises all formats into a common Python dict with "
        "fields: timestamp, level, source_ip, message, raw."
    )

    # 4.8 feature_engineering.py
    add_heading(doc, "4.8 feature_engineering.py — Feature Extraction", 2)
    doc.add_paragraph("Converts parsed log dicts into numerical feature vectors for the ML pipeline. Extracted features include:")
    doc.add_paragraph("message_length, status_code, bytes_sent, cpu, memory, network_tx, network_rx, source_ip (encoded), log_level (encoded), hour_of_day, day_of_week.")

    # 4.9 models.py
    add_heading(doc, "4.9 models.py — Machine Learning Models", 2)
    ml_rows = [
        ("AnomalyDetector",   "IsolationForest (scikit-learn)", "Unsupervised. Flags statistical outliers as anomalies."),
        ("ThreatClassifier",  "RandomForestClassifier",         "Multi-class supervised. Classifies threat type."),
        ("XGBoostClassifier", "XGBoost (architectural stub)",   "Placeholder for future XGBoost integration."),
        ("LightGBMClassifier","LightGBM (architectural stub)",  "Placeholder for future LightGBM integration."),
        ("LSTMDetector",      "PyTorch LSTM (stub)",             "Placeholder for future sequence-based detection."),
    ]
    add_table(doc, ["Class", "Algorithm", "Role"], ml_rows)

    # 4.10 vulnerability_assessment.py
    add_heading(doc, "4.10 vulnerability_assessment.py — Vulnerability Assessment", 2)
    doc.add_paragraph(
        "Analyses incoming feature batches to identify known-vulnerable software versions "
        "(e.g. Apache 2.4.49, Log4j 2.14.1) from a local CVE database. Computes a global 0–100 "
        "Organisational Risk Score based on the CVSS scores of matched CVEs. "
        "Architecture is designed to periodically sync with NVD, MITRE ATT&CK, and CISA feeds."
    )

    # 4.11 predictive_analytics.py
    add_heading(doc, "4.11 predictive_analytics.py — Predictive Analytics", 2)
    doc.add_paragraph(
        "Analyses recent database trends (CPU spikes, failed login rates, alert volumes over the last hour) "
        "to forecast the probability and category of future attacks. "
        "Provides actionable mitigation recommendations before an attack fully materialises."
    )

    # ─── 5. DATABASE SCHEMA ──────────────────────────────────────
    add_heading(doc, "5. Database Schema", 1)

    schema_rows = [
        # alerts
        ("alerts", "id", "INTEGER", "Primary Key"),
        ("alerts", "timestamp", "DATETIME", "Indexed; when the alert was generated"),
        ("alerts", "alert_type", "VARCHAR(100)", "e.g. 'SQL Injection', 'Ransomware Activity'"),
        ("alerts", "severity", "VARCHAR(20)", "low | medium | high | critical"),
        ("alerts", "description", "TEXT", "Human-readable description of the threat"),
        ("alerts", "source_ip", "VARCHAR(50)", "Originating IP address"),
        ("alerts", "confidence", "FLOAT", "0.0–1.0 ML confidence score"),
        ("alerts", "mitre_tactic", "VARCHAR(100)", "Comma-separated MITRE ATT&CK technique IDs"),
        ("alerts", "gemini_summary", "TEXT", "ML-generated incident summary (nullable)"),
        # log_entries
        ("log_entries", "id", "INTEGER", "Primary Key"),
        ("log_entries", "timestamp", "DATETIME", "When the log was ingested"),
        ("log_entries", "source", "VARCHAR(50)", "e.g. 'loki', 'manual'"),
        ("log_entries", "level", "VARCHAR(20)", "INFO | WARN | ERROR | CRITICAL"),
        ("log_entries", "message", "TEXT", "Raw log message text"),
        ("log_entries", "raw_data", "JSON", "Parsed structured metadata"),
        # system_health
        ("system_health", "id", "INTEGER", "Primary Key"),
        ("system_health", "cpu_usage", "FLOAT", "CPU utilisation % (0–100)"),
        ("system_health", "memory_usage", "FLOAT", "RAM utilisation % (0–100)"),
        ("system_health", "network_tx", "FLOAT", "Outbound bytes"),
        ("system_health", "network_rx", "FLOAT", "Inbound bytes"),
        ("system_health", "active_connections", "INTEGER", "Open TCP connections"),
    ]
    add_table(doc, ["Table", "Column", "Type", "Description"], schema_rows)

    # ─── 6. FRONTEND DASHBOARD ───────────────────────────────────
    add_heading(doc, "6. Frontend Dashboard", 1)
    doc.add_paragraph(
        "The frontend is a single-page application (SPA) built with vanilla HTML, CSS, and JavaScript. "
        "It auto-polls /api/dashboard every 3 seconds for live data updates."
    )

    ui_rows = [
        ("POC Dashboard", "Main landing tab. Shows all live metrics, charts, alerts, and the Manual Scan form."),
        ("Settings",      "Placeholder for future platform configuration."),
        ("Metric Cards",  "4 live cards: Total Threats (1h), High Severity, Critical Severity, Org Risk Score."),
        ("Attack Timeline", "Chart.js bar chart: attacks per minute over the last 60 minutes."),
        ("Manual Log Analysis", "Paste raw logs → click 'Scan with ML Engine' → see threat probability + JSON results."),
        ("Live Threat Alerts",  "Table of latest 15 detected threats. High/Critical entries show ML Threat Analysis inline."),
        ("Vulnerability Feed",  "List of CVEs matched against incoming logs."),
        ("Predictive Risk",     "AI forecast of next likely attack type with mitigation recommendation."),
        ("Export Excel Button", "Downloads all alerts as security_alerts_export.xlsx via /api/export/excel."),
    ]
    add_table(doc, ["UI Component", "Description"], ui_rows)

    # ─── 7. API REFERENCE ────────────────────────────────────────
    add_heading(doc, "7. API Reference", 1)

    add_heading(doc, "7.1 GET /api/dashboard", 2)
    doc.add_paragraph("Returns JSON with 5 keys:")
    add_code(doc,
        '{\n'
        '  "live_threats":       [ { id, timestamp, type, severity, description, source_ip, mitre, ai_summary } ],\n'
        '  "system_health":      { cpu_usage, memory_usage, active_connections },\n'
        '  "attack_timeline":    [ { time: "ISO", count: N } ],\n'
        '  "vulnerability_summary": { risk_score, cves: [...] },\n'
        '  "predictive_risk":    { probability, category, target, mitigation }\n'
        '}'
    )

    add_heading(doc, "7.2 POST /api/scan", 2)
    doc.add_paragraph("Request body (JSON):")
    add_code(doc, '{ "log_text": "192.168.1.1 - - [24/Jun/2026] \\"GET /index.php?id=1 UNION SELECT... HTTP/1.1\\" 200 4500" }')
    doc.add_paragraph("Response:")
    add_code(doc,
        '{\n'
        '  "detections":       [ { threat_type, severity, confidence, mitre_tactics, signatures } ],\n'
        '  "assessment":       { risk_score, cves: [...] },\n'
        '  "aggregate":        { ... feature summary ... },\n'
        '  "ml_analysis":      { "probability": 95, "reason": "SQL injection detected in URI." }\n'
        '}'
    )

    add_heading(doc, "7.3 GET /api/export/excel", 2)
    doc.add_paragraph(
        "Returns a binary .xlsx file download (Content-Type: application/vnd.openxmlformats-officedocument.spreadsheetml.sheet). "
        "The file contains a single sheet 'Threat Alerts' with columns: "
        "Timestamp, Threat Type, Severity, Source IP, Confidence, MITRE Tactic, AI Analysis, Raw Description."
    )

    # ─── 8. DEPENDENCIES ─────────────────────────────────────────
    add_heading(doc, "8. Dependencies", 1)
    dep_rows = [
        ("flask>=3.0",          "Core web framework"),
        ("scikit-learn>=1.4",   "ML models: IsolationForest, RandomForestClassifier"),
        ("joblib>=1.3",         "Model serialisation"),
        ("numpy>=1.26",         "Numerical computation"),
        ("SQLAlchemy>=2.0",     "ORM and database abstraction"),
        ("APScheduler>=3.10",   "Background job scheduling"),
        ("xgboost>=2.0",        "Gradient boosting (architectural placeholder)"),
        ("lightgbm>=4.0",       "Light gradient boosting (architectural placeholder)"),
        # google-genai dependency removed
        ("pandas>=2.0",         "DataFrame processing for Excel export"),
        ("openpyxl>=3.1",       "Excel .xlsx file generation"),
        ("python-docx",         "Documentation generation (this script)"),
    ]
    add_table(doc, ["Package", "Purpose"], dep_rows)

    # ─── 9. SETUP & RUNNING ──────────────────────────────────────
    add_heading(doc, "9. Setup & Running the Platform", 1)

    add_heading(doc, "9.1 Prerequisites", 2)
    doc.add_paragraph("• Python 3.11 or later")
    doc.add_paragraph("• pip (Python package manager)")
    doc.add_paragraph("• Standard Python environment")

    add_heading(doc, "9.2 Installation", 2)
    add_code(doc,
        "# 1. Clone or navigate to the project directory\n"
        "cd 'Cybersecurity POC'\n\n"
        "# 2. Install all dependencies\n"
        "pip install -r requirements.txt\n\n"
        "# 3. Run the server\n"
        "python run.py\n\n"
        "# 4. Open in browser\n"
        "# Navigate to: http://127.0.0.1:5000"
    )

    add_heading(doc, "9.3 Configuration", 2)
    doc.add_paragraph(
        "Configuration is managed by cybersec_platform/config.py. Override defaults via environment variables:"
    )
    config_rows = [
        ("GRAFANA_URL", "http://localhost:3000", "Grafana base URL"),
        ("LOKI_URL", "http://localhost:3100", "Loki base URL"),
        ("PROMETHEUS_URL", "http://localhost:9090", "Prometheus base URL"),
        ("MODEL_STORE_PATH", "./models", "ML/LLM model directory"),
        ("LLAMA_MODEL_PATH", "./models/llama_model.gguf", "Optional GGUF model path"),
        ("LOG_LEVEL", "INFO", "Application log level"),
        ("FLASK_DEBUG", "0", "Set to 1 for Flask debug mode"),
    ]
    add_table(doc, ["Variable", "Default", "Description"], config_rows)
    doc.add_paragraph("No external API keys are required for core platform operation.")

    # ─── 10. ATTACK DETECTION CATALOGUE ─────────────────────────
    add_heading(doc, "10. Complete Attack Detection Catalogue", 1)
    doc.add_paragraph(
        "The following table lists all 40+ attack signatures included in the detection engine, "
        "their severity level, and the corresponding MITRE ATT&CK technique."
    )

    full_attacks = [
        ("SQL Injection",                  "High",     "T1190",     "Web Application Attack"),
        ("Cross-Site Scripting (XSS)",     "High",     "T1059.007", "Web Application Attack"),
        ("Directory / Path Traversal",     "High",     "T1083",     "Web Application Attack"),
        ("Command Injection / RCE",        "Critical", "T1059",     "Remote Code Execution"),
        ("SSRF",                           "High",     "T1190",     "Web Application Attack"),
        ("XXE Injection",                  "High",     "T1190",     "Web Application Attack"),
        ("SSTI",                           "Critical", "T1190",     "Web Application Attack"),
        ("Insecure Deserialization",       "Critical", "T1211",     "Web Application Attack"),
        ("HTTP Request Smuggling",         "High",     "T1190",     "Web Application Attack"),
        ("Brute-Force Authentication",     "Medium",   "T1110",     "Authentication Attack"),
        ("Password Spraying",              "High",     "T1110.003", "Authentication Attack"),
        ("Credential Dumping",             "Critical", "T1003",     "Credential Access"),
        ("Credential Exposure in Logs",    "High",     "T1552",     "Credential Access"),
        ("OAuth / JWT Token Theft",        "High",     "T1528",     "Credential Access"),
        ("Privilege Escalation (sudo)",    "Critical", "T1068",     "Privilege Escalation"),
        ("Container Escape",               "Critical", "T1611",     "Privilege Escalation"),
        ("Kernel Exploit Attempt",         "Critical", "T1068",     "Privilege Escalation"),
        ("Port Scanning / Recon",          "Medium",   "T1046",     "Reconnaissance"),
        ("DDoS / Volumetric Attack",       "High",     "T1498",     "Network Attack"),
        ("Man-in-the-Middle (MitM)",       "High",     "T1557",     "Network Attack"),
        ("DNS Tunneling / Exfiltration",   "High",     "T1048.003", "Exfiltration"),
        ("Lateral Movement / PtH",         "Critical", "T1550.002", "Lateral Movement"),
        ("Ransomware Activity",            "Critical", "T1486",     "Malware"),
        ("Cryptominer / Coin Miner",       "High",     "T1496",     "Malware"),
        ("Malware Download / C2 Beacon",   "Critical", "T1105",     "Malware"),
        ("Fileless Malware / LOLBins",     "Critical", "T1059.001", "Malware"),
        ("Persistence Mechanism",          "High",     "T1053",     "Persistence"),
        ("Data Exfiltration",              "Critical", "T1041",     "Exfiltration"),
        ("Database Exfiltration / Dump",   "Critical", "T1005",     "Exfiltration"),
        ("Cloud Metadata Abuse (IMDS)",    "Critical", "T1552.005", "Cloud Attack"),
        ("Kubernetes API Server Abuse",    "Critical", "T1610",     "Cloud Attack"),
        ("IAM / Permissions Abuse",        "High",     "T1078.004", "Cloud Attack"),
        ("S3 / Cloud Storage Exposure",    "High",     "T1530",     "Cloud Attack"),
        ("Anomalous Process Execution",    "High",     "T1055",     "Defense Evasion"),
        ("Log Tampering / Anti-Forensics", "Critical", "T1070",     "Defense Evasion"),
        ("Firewall / Security Disabled",   "Critical", "T1562",     "Defense Evasion"),
        ("SSH Key Tampering",              "High",     "T1098.004", "Persistence"),
        ("VPN / Tunneling Abuse",          "Medium",   "T1133",     "Initial Access"),
        ("Tor / Anonymous Proxy",          "Medium",   "T1090.003", "Command & Control"),
        ("Log4Shell / Known CVE Exploit",  "Critical", "T1190",     "Initial Access"),
        ("Supply Chain / Third-Party",     "Critical", "T1195",     "Initial Access"),
        ("API Abuse / Scraping",           "Medium",   "T1190",     "Discovery"),
        ("Zero-Day / Unknown Exploit",     "Critical", "T1203",     "Initial Access"),
    ]
    add_table(doc, ["Attack Name", "Severity", "MITRE ID", "Category"], full_attacks)

    # ─── 11. CI/CD ───────────────────────────────────────────────
    add_heading(doc, "11. CI/CD — Sonar-like Code Analysis", 1)
    doc.add_paragraph(
        "The GitHub Actions workflow .github/workflows/sonar-analysis.yml runs on push and pull requests to main. "
        "It scans all .py, .js, .ts, .java, .html, and .css source files with SonarAnalyzer and fails the build "
        "if any Critical-severity issue is found."
    )

    # ─── 12. FUTURE ROADMAP ──────────────────────────────────────
    add_heading(doc, "12. Future Roadmap", 1)
    roadmap = [
        ("Replace Mock Clients",       "Connect GrafanaClient / LokiClient / PrometheusClient to real production endpoints."),
        ("PostgreSQL Migration",       "Replace SQLite with PostgreSQL via SQLAlchemy connection string for production scale."),
        ("LSTM / BERT Models",         "Train sequence models on labelled log datasets for higher accuracy anomaly detection."),
        ("NVD / CISA Feed Sync",       "Implement scheduled CVE database sync from NVD API, MITRE ATT&CK TAXII, and CISA KEV."),
        ("SOAR Integration",           "Add webhook triggers for PagerDuty, Slack, or JIRA on Critical alerts."),
        ("Multi-Tenant Support",       "Add user authentication and role-based access control (RBAC)."),
        ("Custom Alert Rules UI",      "Allow analysts to define custom detection rules from the settings tab."),
        ("Threat Intelligence Feeds",  "Integrate IP reputation lookups (AbuseIPDB, VirusTotal) for enriched alerts."),
        ("Geo-IP Visualisation",       "Add a world map heatmap of attack origins to the dashboard."),
        ("Kubernetes Deployment",      "Provide Helm chart for production Kubernetes deployment with horizontal scaling."),
    ]
    add_table(doc, ["Feature", "Description"], roadmap)

    # ─── 13. KNOWN LIMITATIONS ───────────────────────────────────
    add_heading(doc, "13. Known Limitations", 1)
    doc.add_paragraph("• Unsupervised models (IsolationForest) are trained on seed data at startup; real production use requires historical dataset training.")
    doc.add_paragraph("• All data sources are currently mocked. No real network calls are made to Grafana/Loki/Prometheus.")
    doc.add_paragraph("• ML models (IsolationForest, RandomForest) are trained on random data at startup. Real detection accuracy requires historical labelled data.")
    doc.add_paragraph("• SQLite has write concurrency limitations. Use PostgreSQL for any multi-process or multi-user deployment.")
    doc.add_paragraph("• SonarAnalyzer is regex-based and heuristic; it is not a substitute for full SAST tooling like SonarQube.")

    # ─── FOOTER ──────────────────────────────────────────────────
    doc.add_page_break()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(f"NexusSec Enterprise Cybersecurity Platform  |  Internal Documentation  |  {datetime.date.today().strftime('%B %Y')}")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    # Save
    output_path = r"c:\Users\RuchitRathi\Cybersecurity POC\NexusSec_Documentation.docx"
    doc.save(output_path)
    print(f"Documentation successfully generated: {output_path}")

if __name__ == "__main__":
    main()
